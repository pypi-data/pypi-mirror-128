# -*- coding: utf-8 -*-
""" Tableau de marche depuis gpx iPhiGéNie. """

__version__ = "0.1.1"

import argparse
import datetime
import math
import os
import pathlib
import re
import tempfile
import tkinter as tk
import xml.etree.ElementTree as ET
import xml.dom.minidom as MD

from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import Checkbutton
from tkinter import Label
from tkinter import Entry
from tkinter import Menu

import matplotlib.pyplot as plt
import utm

from docx import Document
from docx.enum.section import WD_ORIENT


def get_haversine(lat1, lng1, lat2, lng2):
    """formule de haversine en mètres"""
    _AVG_EARTH_RADIUS_M = 6371008.8
    lat1 = math.radians(lat1)
    lng1 = math.radians(lng1)
    lat2 = math.radians(lat2)
    lng2 = math.radians(lng2)
    lat = lat2 - lat1
    lng = lng2 - lng1
    d = math.sin(lat * 0.5) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin(lng * 0.5) ** 2
    return 2 * _AVG_EARTH_RADIUS_M * math.asin(math.sqrt(d))


def get_azimuth(lat1, lon1, lat2, lon2):
    """Calcul azimuth"""
    dlon = math.radians(lon2 - lon1)
    lat1 = math.radians(lat1)
    lat2 = math.radians(lat2)
    x = math.cos(lat2) * math.sin(dlon)
    y = math.cos(lat1) * math.sin(lat2) - math.sin(lat1) * math.cos(lat2) * math.cos(dlon)
    azimut = round(((math.atan2(x, y) * 180 / math.pi) + 360) % 360)
    return azimut


def prettify_xml(elem):
    """Return a pretty-printed XML string for the Element."""
    rough_string = ET.tostring(elem, "utf-8")
    reparsed = MD.parseString(rough_string)
    return reparsed.toprettyxml(indent="  ")


def get_parser():
    ap = argparse.ArgumentParser(
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    ap.add_argument("-f", "--gpx-file", type=pathlib.Path, required=False, help="Fichier à traiter")
    ap.add_argument("-V", "--version", action="version", version=__version__)
    ap.add_argument(
        "--retour", action=argparse.BooleanOptionalAction, default=False, help="Génère le retour"
    )
    ap.add_argument(
        "--gui", action=argparse.BooleanOptionalAction, default=True, help="Mode graphique"
    )
    subparsers = ap.add_subparsers(title="subcommands", dest="subcmd")

    # ipgn2tdm wpt
    parser_wpt = subparsers.add_parser(
        "wpt",
        help="Créer un fichier wpt",
    )
    parser_wpt.add_argument(
        "-w",
        "--wpt-file",
        type=pathlib.Path,
        default="wpt.gpx",
        help="Nom de fichier de sortie (défaut wpt.gpx)",
    )
    # ipgn2tdm tdm
    parser_tdm = subparsers.add_parser(
        "tdm",
        help="Calculer un tableau de marche",
    )
    parser_tdm.add_argument(
        "--plat",
        type=int,
        default=4000,
        help="distance m/h à plat (défaut 4000)",
    )
    parser_tdm.add_argument(
        "--asc",
        type=int,
        default=300,
        help="dénivelé m/h en montée (défaut 300)",
    )
    parser_tdm.add_argument(
        "--des",
        type=int,
        default=450,
        help="dénivelé m/h en descente (défaut 450)",
    )
    parser_tdm.add_argument(
        "--pause",
        type=int,
        default=10,
        help="temps de pause en %% (défaut 10)",
    )
    parser_tdm.add_argument(
        "--profil",
        action=argparse.BooleanOptionalAction,
        default=False,
        help="Génère le profil d'altitude",
    )
    parser_tdm.add_argument(
        "--docx-file",
        type=pathlib.Path,
        default="tdm.docx",
        help="Nom de fichier docx de sortie (défaut tdm.docx)",
    )
    return ap


def wpt(root, output_file, retour):
    """Création d'un fichier avec points de passage (wpt) pour https://istresrando.fr/gpxRando/"""
    data = ET.Element("gpx")
    trk = ET.SubElement(data, "trk")
    trkseg = ET.SubElement(trk, "trkseg")
    ns = re.match(r"{.*}", root.tag).group(0)
    trk = root.find(f"{ns}trk/{ns}trkseg")
    for i in trk:
        trkpt = ET.SubElement(trkseg, "trkpt")
        e = i.find(f"{ns}ele")
        h = i.find(f"{ns}time")
        trkpt.set("lat", i.attrib["lat"])
        trkpt.set("lon", i.attrib["lon"])
        ele = ET.SubElement(trkpt, "ele")
        ele.text = e.text
        twpt = ET.SubElement(trkpt, "time")
        twpt.text = h.text
        t = i.find(f"{ns}type")
        if isinstance(t, ET.Element) and t.text == "iPGN_wpt":
            if isinstance(i.find(f"{ns}link/{ns}text"), ET.Element):
                wpt = ET.SubElement(data, "wpt")
                name = ET.SubElement(wpt, "name")
                name.text = i.find(f"{ns}link/{ns}text").text
                wpt.set("lat", i.attrib["lat"])
                wpt.set("lon", i.attrib["lon"])
                ele = ET.SubElement(wpt, "ele")
                ele.text = e.text
                twpt = ET.SubElement(wpt, "time")
                twpt.text = h.text
    if retour:
        for i in reversed(trk):
            trkpt = ET.SubElement(trkseg, "trkpt")
            e = i.find(f"{ns}ele")
            h = i.find(f"{ns}time")
            trkpt.set("lat", i.attrib["lat"])
            trkpt.set("lon", i.attrib["lon"])
            ele = ET.SubElement(trkpt, "ele")
            ele.text = e.text
            twpt = ET.SubElement(trkpt, "time")
            twpt.text = h.text
    with open(output_file, "w") as f:
        f.write(prettify_xml(data))


def tdm(root, retour, vitesse, montee, descente, pause):
    """Générer un tableau de marche"""
    ns = re.match(r"{.*}", root.tag).group(0)
    trk = root.find(f"{ns}trk/{ns}trkseg")
    first = True
    tdm_dict = dict()
    tdm_list = []
    ele_list = []
    distance_list = []
    for i in trk:
        e = i.find(f"{ns}ele")
        if first:
            first = False
            utm_val = utm.from_latlon(float(i.attrib["lat"]), float(i.attrib["lon"]))
            distance = 0
            cumul_distance = 0
            positif = 0
            negatif = 0
            lat_ref = float(i.attrib["lat"])
            lon_ref = float(i.attrib["lon"])
            ele_ref = int(e.text)
            ele_list.append(int(e.text))
            total_distance = 0 + distance
            distance_list.append(total_distance)
            azimut = None
            temps_segment = 0
            cumul_temps = 0
            tdm_dict["nom"] = "Départ"
            tdm_dict["utm"] = f"{utm_val[2]}{utm_val[3]} {round(utm_val[0])} {round(utm_val[1])}"
            tdm_dict["azimut"] = ""
            tdm_dict["altitude"] = ele_ref
            tdm_dict["positif"] = positif
            tdm_dict["négatif"] = negatif
            tdm_dict["distance"] = 0
            tdm_dict["cumul distance"] = 0
            tdm_dict["temps segment"] = "0"
            tdm_dict["temps cumulé"] = "0"
            tdm_list.append(tdm_dict)
        else:
            lat = float(i.attrib["lat"])
            lon = float(i.attrib["lon"])
            partiel = get_haversine(lat_ref, lon_ref, lat, lon)
            ele_list.append(int(e.text))
            total_distance = total_distance + partiel
            distance_list.append(total_distance)
            distance = distance + partiel
            if not azimut and lat != lat_ref and lon != lon_ref:
                azimut = round(get_azimuth(lat_ref, lon_ref, lat, lon))
            lat_ref = lat
            lon_ref = lon
            t = i.find(f"{ns}type")
            ele = int(e.text)
            if (ele - ele_ref) > 0:
                positif = positif + (ele - ele_ref)
            else:
                negatif = negatif + (ele_ref - ele)
            try:
                pente = ((ele - ele_ref) / partiel) * 100
            except ZeroDivisionError:
                pente = 0
            if pente > 10:
                temps = 60 / montee * (ele - ele_ref)
            elif pente < -10:
                temps = 60 / descente * (ele_ref - ele)
            else:
                temps = 60 / vitesse * partiel
            temps_segment = temps_segment + temps
            ele_ref = ele
            if (isinstance(t, ET.Element) and t.text == "iPGN_wpt") or i == trk[-1]:
                tdm_dict = dict()
                if isinstance(i.find(f"{ns}link/{ns}text"), ET.Element):
                    tdm_dict["nom"] = i.find(f"{ns}link/{ns}text").text
                else:
                    tdm_dict["nom"] = "Arrivée"
                utm_val = utm.from_latlon(lat, lon)
                tdm_dict["utm"] = f"{utm_val[2]}{utm_val[3]} {int(utm_val[0])} {int(utm_val[1])}"
                tdm_dict["azimut"] = azimut
                tdm_dict["altitude"] = ele
                tdm_dict["positif"] = positif
                tdm_dict["négatif"] = negatif
                tdm_dict["distance"] = round(distance)
                cumul_distance = cumul_distance + distance
                tdm_dict["cumul distance"] = round(cumul_distance)
                temps_segment = temps_segment * (1 + pause / 100)
                tdm_dict["temps segment"] = str(datetime.timedelta(minutes=round(temps_segment)))[
                    :-3
                ]
                cumul_temps = cumul_temps + temps_segment
                tdm_dict["temps cumulé"] = str(datetime.timedelta(minutes=round(cumul_temps)))[:-3]
                tdm_list.append(tdm_dict)
                distance = 0
                temps_segment = 0
                positif = 0
                negatif = 0
                azimut = None
    if retour:
        first = True
        for i in reversed(trk):
            e = i.find(f"{ns}ele")
            if first:
                first = False
            else:
                lat = float(i.attrib["lat"])
                lon = float(i.attrib["lon"])
                partiel = get_haversine(lat_ref, lon_ref, lat, lon)
                ele_list.append(int(e.text))
                total_distance = total_distance + partiel
                distance_list.append(total_distance)
                distance = distance + partiel
                if not azimut and lat != lat_ref and lon != lon_ref:
                    azimut = round(get_azimuth(lat_ref, lon_ref, lat, lon))
                lat_ref = lat
                lon_ref = lon
                t = i.find(f"{ns}type")
                ele = int(e.text)
                if (ele - ele_ref) > 0:
                    positif = positif + (ele - ele_ref)
                else:
                    negatif = negatif + (ele_ref - ele)
                try:
                    pente = ((ele - ele_ref) / partiel) * 100
                except ZeroDivisionError:
                    pente = 0
                if pente > 10:
                    temps = 60 / montee * (ele - ele_ref)
                elif pente < -10:
                    temps = 60 / descente * (ele_ref - ele)
                else:
                    temps = 60 / vitesse * partiel
                temps_segment = temps_segment + temps
                ele_ref = ele
                if (isinstance(t, ET.Element) and t.text == "iPGN_wpt") or i == trk[0]:
                    tdm_dict = dict()
                    if isinstance(i.find(f"{ns}link/{ns}text"), ET.Element):
                        tdm_dict["nom"] = i.find(f"{ns}link/{ns}text").text
                    else:
                        tdm_dict["nom"] = "Arrivée"
                    utm_val = utm.from_latlon(lat, lon)
                    tdm_dict[
                        "utm"
                    ] = f"{utm_val[2]}{utm_val[3]} {int(utm_val[0])} {int(utm_val[1])}"
                    tdm_dict["azimut"] = azimut
                    tdm_dict["altitude"] = ele
                    tdm_dict["positif"] = positif
                    tdm_dict["négatif"] = negatif
                    tdm_dict["distance"] = round(distance)
                    cumul_distance = cumul_distance + distance
                    tdm_dict["cumul distance"] = round(cumul_distance)
                    temps_segment = temps_segment * (1 + pause / 100)
                    tdm_dict["temps segment"] = str(
                        datetime.timedelta(minutes=round(temps_segment))
                    )[:-3]
                    cumul_temps = cumul_temps + temps_segment
                    tdm_dict["temps cumulé"] = str(datetime.timedelta(minutes=round(cumul_temps)))[
                        :-3
                    ]
                    tdm_list.append(tdm_dict)
                    distance = 0
                    temps_segment = 0
                    positif = 0
                    negatif = 0
                    azimut = None
    return tdm_list, ele_list, distance_list


def create_docx(
    to_csv, ele_list, distance_list, profil, docx_file, nom, carte, date, ibp, cotation
):
    """Créer un tableau de marche docx"""
    document = Document()
    sections = document.sections
    for section in sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
    document.add_heading("Tableau de marche", 0)
    document.add_paragraph(f"NOM de la RANDONNEE : {nom} Date: {date}")
    document.add_paragraph(f"CARTE IGN : {carte}")
    document.add_paragraph(f"IBP Index : {ibp} Cotation: {cotation}")
    nb_cols = len(to_csv[0].keys()) + 1
    table = document.add_table(rows=1, cols=nb_cols, style="Table Grid")
    table.autofit = True
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    for i, j in enumerate(to_csv[0].keys()):
        hdr_cells[i].text = j
    hdr_cells[nb_cols - 1].text = "Notes et observations"
    t_positif = 0
    t_negatif = 0
    t_distance = 0
    for i in to_csv:
        row_cells = table.add_row().cells
        t_positif = t_positif + i["positif"]
        t_negatif = t_negatif + i["négatif"]
        t_distance = t_distance + i["distance"]
        for j, z in enumerate(i.values()):
            row_cells[j].text = str(z)
    row_cells = table.add_row().cells
    row_cells[0].text = "Total"
    row_cells[4].text = str(t_positif)
    row_cells[5].text = str(t_negatif)
    row_cells[6].text = str(t_distance)
    if profil:
        plt.plot(distance_list, ele_list)
        plt.xlabel("Distance en m")
        plt.ylabel("Altitude en m")
        plt.title("Profil altitude")
        plt.axis("scaled")
        plt.grid()
        png_file = os.path.join(tempfile.gettempdir(), "profil.png")
        plt.savefig(png_file, bbox_inches="tight")
        document.add_paragraph("")
        document.add_picture(png_file)
    document.save(docx_file)


def gui():
    global gpx_file
    gpx_file = None

    def _open_gpx():
        global gpx_file
        gpx_file = filedialog.askopenfilename(
            initialdir=pathlib.Path.home(),
            filetypes=(("Fichier gpx", "*.gpx"), ("all files", "*.*")),
            title="Choisir fichier gpx (iPhiGéNie)",
        )

    def _wpt():
        if not gpx_file:
            _open_gpx()
        wpt_file = filedialog.asksaveasfilename(
            initialdir=pathlib.Path.home(),
            filetypes=(("Fichier gpx", "*.gpx"), ("all files", "*.*")),
            initialfile="wpt.gpx",
            title="Fichier gpx avec point de passage",
        )
        tree = ET.parse(gpx_file)
        root = tree.getroot()
        wpt(root, wpt_file, retour_state.get())

    def _tdm():
        if not gpx_file:
            _open_gpx()
        tdm_file = filedialog.asksaveasfilename(
            initialdir=pathlib.Path.home(),
            filetypes=(("Fichier tdm", "*.docx"), ("all files", "*.*")),
            initialfile="tdm.docx",
            title="Fichier tableau de marche",
        )
        tree = ET.parse(gpx_file)
        root = tree.getroot()
        to_csv, ele_list, distance_list = tdm(
            root,
            retour_state.get(),
            int(plat_txt.get()),
            int(asc_txt.get()),
            int(des_txt.get()),
            int(repos_txt.get()),
        )
        create_docx(
            to_csv,
            ele_list,
            distance_list,
            profil_state.get(),
            tdm_file,
            nom_txt.get(),
            carte_txt.get(),
            date_txt.get(),
            ibp_txt.get(),
            cotation_txt.get(),
        )

    def _about():
        apropos = "Licence BSD 3-Clause \nCopyright (c) 2021, Philippe Makowski \nUtilise les modules utm, matplotlib et docx."
        messagebox.showinfo(titre, message=apropos)

    root = tk.Tk()
    titre = "Tableau de marche depuis gpx iPhiGéNie."
    root.title(titre)
    titre_lbl = Label(root, text=titre)
    titre_lbl.grid(column=0, row=0)
    carte_lbl = Label(
        root,
        text="Pour imprimer la carte utilisez le fichier wpt avec le site \n https://istresrando.fr/gpxRando/",
    )
    carte_lbl.grid(column=0, row=1)

    frame = ttk.Frame(root)
    options = {"padx": 5, "pady": 5}
    gpx_file_button = ttk.Button(frame, text="Choisir fichier gpx (iPhiGéNie)")
    gpx_file_button.grid(column=0, row=0, sticky="W", **options)
    gpx_file_button.configure(command=_open_gpx)
    retour_state = tk.BooleanVar()
    retour_state.set(False)
    chk_retour = Checkbutton(frame, text="aller retour", var=retour_state)
    chk_retour.grid(column=0, row=1)
    wpt_button = ttk.Button(frame, text="Créer fichier wpt")
    wpt_button.grid(column=1, row=1, sticky="W", **options)
    wpt_button.configure(command=_wpt)

    # nom, carte, date, ibp, cotation
    nom_lbl = Label(frame, text="Randonnée : ")
    nom_lbl.grid(column=0, row=2)
    nom_txt = Entry(frame, width=20)
    nom_txt.grid(column=1, row=2)
    carte_lbl = Label(frame, text="Carte : ")
    carte_lbl.grid(column=0, row=3)
    carte_txt = Entry(frame, width=20)
    carte_txt.grid(column=1, row=3)
    date_lbl = Label(frame, text="Date : ")
    date_lbl.grid(column=0, row=4)
    date_txt = Entry(frame, width=20)
    date_txt.grid(column=1, row=4)
    ibp_lbl = Label(frame, text="IBP index : ")
    ibp_lbl.grid(column=0, row=5)
    ibp_txt = Entry(frame, width=20)
    ibp_txt.grid(column=1, row=5)
    cotation_lbl = Label(frame, text="Cotation : ")
    cotation_lbl.grid(column=0, row=6)
    cotation_txt = Entry(frame, width=20)
    cotation_txt.grid(column=1, row=6)
    plat_lbl = Label(frame, text="Vitesse à plat m/h : ")
    plat_lbl.grid(column=0, row=7)
    plat_txt = Entry(frame, width=20)
    plat_txt.grid(column=1, row=7)
    plat_txt.delete(0, tk.END)
    plat_txt.insert(0, "4000")
    asc_lbl = Label(frame, text="Vitesse montée m/h : ")
    asc_lbl.grid(column=0, row=8)
    asc_txt = Entry(frame, width=20)
    asc_txt.grid(column=1, row=8)
    asc_txt.delete(0, tk.END)
    asc_txt.insert(0, "300")
    des_lbl = Label(frame, text="Vitesse descente m/h : ")
    des_lbl.grid(column=0, row=9)
    des_txt = Entry(frame, width=20)
    des_txt.grid(column=1, row=9)
    des_txt.delete(0, tk.END)
    des_txt.insert(0, "450")
    repos_lbl = Label(frame, text="Temps de repos en % : ")
    repos_lbl.grid(column=0, row=10)
    repos_txt = Entry(frame, width=20)
    repos_txt.grid(column=1, row=10)
    repos_txt.delete(0, tk.END)
    repos_txt.insert(0, "10")

    profil_state = tk.BooleanVar()
    profil_state.set(False)
    profil_retour = Checkbutton(frame, text="Profil altitude", var=profil_state)
    profil_retour.grid(column=0, row=11)
    tdm_button = ttk.Button(frame, text="Créer fichier tdm")
    tdm_button.grid(column=1, row=11, sticky="W", **options)
    tdm_button.configure(command=_tdm)

    exit_button = ttk.Button(frame, text="Quitter", command=root.destroy)
    exit_button.grid(column=0, row=12, sticky="W", **options)

    frame.grid(padx=10, pady=10)
    root.option_add("*tearOff", tk.FALSE)
    menubar = Menu(root)
    filemenu = Menu(menubar)
    filemenu.add_command(label="Fichier wpt", command=_wpt)
    filemenu.add_command(label="Fichier tdm", command=_tdm)
    filemenu.add_separator()
    filemenu.add_command(label="Quitter", command=root.quit)
    menubar.add_cascade(label="Fichier", menu=filemenu)
    menu_help = Menu(menubar, name="help")
    menu_help.add_command(label="À propos...", command=_about)
    menubar.add_cascade(label="Aide", menu=menu_help)
    root.config(menu=menubar)
    root.mainloop()


def no_gui():
    parser = get_parser()
    args = parser.parse_args()
    if not args.gpx_file:
        print("Error : the following arguments are required: -f/--gpx-file")
        exit(1)
    if args.subcmd == "wpt":
        tree = ET.parse(args.gpx_file)
        root = tree.getroot()
        wpt(root, args.wpt_file, args.retour)
    elif args.subcmd == "tdm":
        tree = ET.parse(args.gpx_file)
        root = tree.getroot()
        to_csv, ele_list, distance_list = tdm(
            root, args.retour, args.plat, args.asc, args.des, args.pause
        )
        create_docx(
            to_csv, ele_list, distance_list, args.profil, args.docx_file, "", "", "", "", ""
        )


def main():
    parser = get_parser()
    args = parser.parse_args()
    if args.gui:
        gui()
    else:
        no_gui()


if __name__ == "__main__":
    main()
